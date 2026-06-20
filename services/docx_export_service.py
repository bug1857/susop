import os
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

class DocxExportService:
    @staticmethod
    def generate_docx(report: Dict[str, Any], output_path: str) -> str:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document()
        
        # Configure margins: 1 inch on all sides
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Style helpers
        def set_font(run, name="Arial", size=10, bold=False, color=None):
            run.font.name = name
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = color

        # Colors
        navy = RGBColor(30, 58, 138)
        blue = RGBColor(37, 99, 235)
        charcoal = RGBColor(55, 65, 81)
        gray = RGBColor(107, 114, 128)
        green = RGBColor(22, 163, 74)

        # ------------------ COVER PAGE ------------------
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(80)

        p_title = doc.add_paragraph()
        run_title = p_title.add_run("SustainOCPM Compliance Report")
        set_font(run_title, size=26, bold=True, color=navy)
        
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Business Responsibility and Sustainability Reporting (BRSR) Disclosures")
        set_font(run_sub, size=14, color=gray)
        p_sub.paragraph_format.space_after = Pt(40)

        # Divider line
        p_line = doc.add_paragraph()
        p_line.add_run("__________________________________________________________________").font.color.rgb = blue
        p_line.paragraph_format.space_after = Pt(40)

        # Metadata Box (Table)
        sec_a = report.get("section_a", {})
        meta_items = [
            ("Organization Name:", str(sec_a.get("organization_name", "N/A"))),
            ("Workspace Context:", str(sec_a.get("workspace_name", "N/A"))),
            ("Project Context:", str(sec_a.get("project_name", "N/A"))),
            ("Reporting Period:", str(sec_a.get("reporting_period", "N/A"))),
            ("Report Version:", f"Version {report.get('report_version', 1)}"),
            ("Audit Readiness:", str(report.get("audit_readiness", "N/A"))),
            ("Report SHA256 Hash:", str(report.get("sha256_hash", "N/A"))),
        ]
        
        table_meta = doc.add_table(rows=len(meta_items), cols=2)
        table_meta.autofit = False
        table_meta.columns[0].width = Inches(2.2)
        table_meta.columns[1].width = Inches(4.3)
        
        for idx, (label, val) in enumerate(meta_items):
            row = table_meta.rows[idx]
            
            p_lbl = row.cells[0].paragraphs[0]
            run_lbl = p_lbl.add_run(label)
            set_font(run_lbl, size=10, bold=True, color=charcoal)
            
            p_val = row.cells[1].paragraphs[0]
            run_val = p_val.add_run(val)
            if label == "Report SHA256 Hash:":
                set_font(run_val, name="Courier New", size=8, color=charcoal)
            else:
                set_font(run_val, size=10, color=charcoal)

        doc.add_page_break()

        # Helper for Heading 1
        def add_h1(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            set_font(run, size=18, bold=True, color=navy)
            return p

        # Helper for Heading 2
        def add_h2(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_font(run, size=12, bold=True, color=blue)
            return p

        # Helper for Body
        def add_body(text, bold=False, italic=False, space_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(space_after)
            run = p.add_run(text)
            set_font(run, size=10, bold=bold, color=charcoal)
            run.italic = italic
            return p

        # ------------------ EXECUTIVE OVERVIEW ------------------
        add_h1("Executive Summary")
        add_body(report.get("executive_summary", "Summary not generated."))

        add_h2("Key Compliance & Sustainability KPIs")
        sec_b = report.get("section_b", {})
        sec_c = report.get("section_c", {})
        
        kpis = [
            ("Process Compliance Score", f"{sec_b.get('compliance_score', 0.0) * 100:.1f}%"),
            ("Carbon Fitness Score", f"{sec_b.get('carbon_fitness', 0.0) * 100:.1f}%"),
            ("ESG Overall Score", f"{sec_c.get('esg_overall_score', 0.0) * 100:.1f}%"),
            ("Total Actual Emissions", f"{sec_c.get('total_actual_emissions_kg', 0.0):,.1f} kg")
        ]
        
        table_kpi = doc.add_table(rows=2, cols=4)
        for c_idx, (kpi_name, kpi_val) in enumerate(kpis):
            # Header
            cell_hdr = table_kpi.rows[0].cells[c_idx]
            cell_hdr.text = kpi_name
            set_font(cell_hdr.paragraphs[0].runs[0], size=9, bold=True, color=navy)
            
            # Value
            cell_val = table_kpi.rows[1].cells[c_idx]
            cell_val.text = kpi_val
            set_font(cell_val.paragraphs[0].runs[0], size=11, bold=True, color=blue)
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(10)

        # ------------------ SECTION A ------------------
        add_h1("Section A: General Disclosures")
        add_body("This section outlines basic details of the reporting organization and scope boundary.", italic=True)
        
        table_sec_a = doc.add_table(rows=len(meta_items) - 1, cols=2)
        for idx, (label, val) in enumerate(meta_items[:-1]):
            row = table_sec_a.rows[idx]
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].paragraphs[0].text = val
            set_font(row.cells[0].paragraphs[0].runs[0], size=9, color=charcoal)
            set_font(row.cells[1].paragraphs[0].runs[0], size=9, color=charcoal)

        # ------------------ SECTION B ------------------
        add_h1("Section B: Process & Management Disclosures")
        add_body(f"Conformance Check Methodology: {sec_b.get('conformance_method', 'N/A')}", bold=True)
        add_body(f"Total Evaluated Traces: {sec_b.get('total_trace_count', 0):,}")
        add_body(f"Non-Conforming Traces: {sec_b.get('non_conforming_trace_count', 0):,}")
        
        bots = sec_b.get("bottlenecks", [])
        if bots:
            add_h2("Significant Bottlenecks Impacting Wait Times")
            table_bots = doc.add_table(rows=1 + len(bots), cols=3)
            # Headers
            table_bots.rows[0].cells[0].text = "Activity Name"
            table_bots.rows[0].cells[1].text = "Average Wait Time (sec)"
            table_bots.rows[0].cells[2].text = "Occurrences"
            for cell in table_bots.rows[0].cells:
                set_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=navy)
                
            for idx, b in enumerate(bots):
                row = table_bots.rows[idx + 1]
                row.cells[0].text = b.get("activity_name", "")
                row.cells[1].text = f"{b.get('average_wait_time_sec', 0.0):,.1f}"
                row.cells[2].text = str(b.get("occurrence_count", 0))
                for cell in row.cells:
                    set_font(cell.paragraphs[0].runs[0], size=9, color=charcoal)

        # ------------------ SECTION C ------------------
        add_h1("Section C: Principle-wise Performance Disclosures")
        add_h2("Resource Draw & Footprint Mappings")
        
        res_items = [
            ("Energy Consumption (kWh)", f"{sec_c.get('total_energy_consumption_kwh', 0.0):,.1f} kWh"),
            ("Water Consumption (Liters)", f"{sec_c.get('total_water_consumption_liters', 0.0):,.1f} L"),
            ("Waste Generated (kg)", f"{sec_c.get('total_waste_generation_kg', 0.0):,.1f} kg"),
            ("Carbon Budget Limit", f"{sec_c.get('carbon_budget_limit_kg', 0.0):,.1f} kg CO2e"),
            ("Carbon Budget Status", "EXCEEDED" if sec_c.get("carbon_budget_exceeded") else "CONFORMING"),
        ]
        
        table_res = doc.add_table(rows=len(res_items), cols=2)
        for idx, (label, val) in enumerate(res_items):
            row = table_res.rows[idx]
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].paragraphs[0].text = val
            set_font(row.cells[0].paragraphs[0].runs[0], size=9, color=charcoal)
            if val in ["EXCEEDED", "CONFORMING"]:
                color = RGBColor(220, 38, 38) if val == "EXCEEDED" else green
                set_font(row.cells[1].paragraphs[0].runs[0], size=9, bold=True, color=color)
            else:
                set_font(row.cells[1].paragraphs[0].runs[0], size=9, color=charcoal)

        hotspots = sec_c.get("carbon_hotspots", [])
        if hotspots:
            add_h2("High-Impact Carbon Hotspots")
            table_hots = doc.add_table(rows=1 + len(hotspots), cols=3)
            table_hots.rows[0].cells[0].text = "Activity Name"
            table_hots.rows[0].cells[1].text = "Emissions (kg CO2e)"
            table_hots.rows[0].cells[2].text = "Contribution (%)"
            for cell in table_hots.rows[0].cells:
                set_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=navy)
                
            for idx, h in enumerate(hotspots):
                row = table_hots.rows[idx + 1]
                row.cells[0].text = h.get("activity_name", "")
                row.cells[1].text = f"{h.get('emissions_kg', 0.0):,.1f}"
                row.cells[2].text = f"{h.get('contribution_percentage', 0.0):.1f}%"
                for cell in row.cells:
                    set_font(cell.paragraphs[0].runs[0], size=9, color=charcoal)

        # ------------------ SECTION D ------------------
        add_h1("Section D: Traceability Matrix")
        trace = report.get("section_d", {}).get("traceability_matrix", [])
        if trace:
            table_trace = doc.add_table(rows=1 + len(trace), cols=4)
            table_trace.rows[0].cells[0].text = "BRSR Metric"
            table_trace.rows[0].cells[1].text = "Originating Engine"
            table_trace.rows[0].cells[2].text = "Database Table / Source"
            table_trace.rows[0].cells[3].text = "Reference Field"
            for cell in table_trace.rows[0].cells:
                set_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=navy)
                
            for idx, t in enumerate(trace):
                row = table_trace.rows[idx + 1]
                row.cells[0].text = t.get("brsr_metric", "")
                row.cells[1].text = t.get("source_engine", "")
                row.cells[2].text = t.get("database_source", "")
                row.cells[3].text = t.get("reference_field", "")
                for cell in row.cells:
                    set_font(cell.paragraphs[0].runs[0], size=9, color=charcoal)

        # ------------------ RECOMMENDATIONS ------------------
        recs = report.get("recommendations", [])
        if recs:
            doc.add_page_break()
            add_h1("Compliance & Sustainability Recommendations")
            add_body("These recommendations have been dynamically derived based on conformance, bottleneck, carbon hotspot, and supplier risk metrics.", italic=True)
            
            for r in recs:
                add_h2(f"{r.get('title', '')} [Priority: {r.get('priority', '')}]")
                add_body(str(r.get('description', '')))
                
                savings_text = []
                if r.get('estimated_emission_reduction', 0) > 0:
                    savings_text.append(f"Est. Emission Reduction: {r.get('estimated_emission_reduction'):,.1f} kg CO2e")
                if r.get('estimated_cost_reduction', 0) > 0:
                    savings_text.append(f"Est. Cost Savings: ${r.get('estimated_cost_reduction'):,.2f}")
                
                if savings_text:
                    p_savings = doc.add_paragraph()
                    p_savings.paragraph_format.space_after = Pt(12)
                    run_sav = p_savings.add_run(" | ".join(savings_text))
                    set_font(run_sav, size=9, bold=True, color=green)

        doc.save(output_path)
        return output_path
